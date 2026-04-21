#!/usr/bin/env python3
"""
Standalone Batch Processor — Single File, No Subprocess

Merges process_all_lite.py (orchestrator) and process_single_file.py (processor)
into one self-contained file. Eliminates subprocess overhead.

Uses ONLY the configured model with no fallback.
On any error, waits and retries from unprocessed files.
Loops until every file is processed.

Usage:
    python process_standalone.py --course "Microeconomics"
    python process_standalone.py --course "Microeconomics" --images
    python process_standalone.py --course "Microeconomics" --wait 15
"""

import os
import sys
import json
import re
import time
from pathlib import Path
from datetime import datetime, timezone
from dotenv import load_dotenv
import google.generativeai as genai
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

try:
    import pymupdf
except ImportError:
    pymupdf = None

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

PROJECT_ROOT = Path(__file__).resolve().parent.parent
# Ensure project root is on sys.path for cross-package imports
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / '.env')

# ── Configuration ─────────────────────────────────────────────────────────────
GEMINI_API_KEY = os.getenv('Gemini_Api_Key')
MODEL = os.getenv("GEMINI_MODEL", "gemini-3.1-flash-lite-preview")
GEMINI_MODELS = [MODEL]  # Single model, no fallback

WIKI_DIR = PROJECT_ROOT / 'MBAWiki'
CONCEPT_FILE_PREFIX = "Concept-"
CASE_FILE_PREFIX = "Case-"
CONCEPT_FILE_SUFFIX = ".md"

COURSES_FILE = PROJECT_ROOT / 'courses.json'
COURSE_GROUPS_FILE = PROJECT_ROOT / 'course_groups.json'
TRACKER_FILE = PROJECT_ROOT / 'processed_files.json'
LOCAL_DIR = PROJECT_ROOT / 'Transcript_class_lecture'
LOG_FILE = PROJECT_ROOT / 'log.md'
DEFAULT_WAIT_MINUTES = 30

_gemini_configured = False


# ══════════════════════════════════════════════════════════════════════════════
#  JSON Repair Helper
# ══════════════════════════════════════════════════════════════════════════════

def repair_json(text):
    """Try to parse JSON, repairing common LLM issues if needed."""
    # 1. Try as-is
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 2. Fix unescaped newlines/tabs inside strings
    fixed = re.sub(r'(?<=": ")(.*?)(?="[,\}\]])', lambda m: m.group(0).replace('\n', '\\n').replace('\t', '\\t'), text, flags=re.DOTALL)
    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # 3. Fix unescaped newlines inside arrays of strings
    fixed = re.sub(r'(?<=\[)(.*?)(?=\])', lambda m: m.group(0).replace('\n', '\\n').replace('\t', '\\t'), text, flags=re.DOTALL)
    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # 4. Brute force: escape all control chars inside JSON string values
    def escape_control_chars(s):
        out = []
        in_string = False
        escape = False
        for ch in s:
            if escape:
                out.append(ch)
                escape = False
                continue
            if ch == '\\' and in_string:
                out.append(ch)
                escape = True
                continue
            if ch == '"':
                in_string = not in_string
                out.append(ch)
                continue
            if in_string and ch in ('\n', '\r', '\t'):
                out.append('\\n' if ch == '\n' else '\\t' if ch == '\t' else '\\r')
                continue
            out.append(ch)
        return ''.join(out)

    try:
        return json.loads(escape_control_chars(text))
    except json.JSONDecodeError:
        pass

    # 5. Truncation recovery
    last_complete = text.rfind('},')
    if last_complete > 0:
        truncated = text[:last_complete + 1]
        for closing in [']}\n', ']\n}', ']}']:
            try:
                return json.loads(truncated + closing)
            except json.JSONDecodeError:
                continue
        try:
            return json.loads(escape_control_chars(truncated) + ']}')
        except json.JSONDecodeError:
            pass

    return None


# ══════════════════════════════════════════════════════════════════════════════
#  Logging
# ══════════════════════════════════════════════════════════════════════════════

def log_ingestion(action, source_filename, file_type, details=""):
    """Append an entry to log.md with timestamp."""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry = f"## [{timestamp}] {action} | {file_type.capitalize()}: {source_filename}"
        if details:
            entry += f" | {details}"
        entry += "\n"

        if LOG_FILE.exists():
            with open(LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(entry)
        else:
            header = "# Wiki Evolution Log\n\nAppend-only record of ingestions and updates.\n\n"
            with open(LOG_FILE, 'w', encoding='utf-8') as f:
                f.write(header)
                f.write(entry)
    except Exception as e:
        print(f"   Warning: Could not log entry: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  Course Groups & Concept Loading
# ══════════════════════════════════════════════════════════════════════════════

def load_course_groups(course_name):
    """Load course_groups.json and return sibling course names for the given course."""
    if not COURSE_GROUPS_FILE.exists():
        return None

    try:
        with open(COURSE_GROUPS_FILE, 'r', encoding='utf-8') as f:
            groups = json.load(f)

        siblings = set()
        for group_name, courses in groups.items():
            if course_name in courses:
                siblings.update(c for c in courses if c != course_name)

        return sorted(siblings) if siblings else None
    except Exception as e:
        print(f"   Warning: Could not load course_groups.json: {e}")

    return None


def _extract_course_from_file(filepath):
    """Extract course names from the first 5 lines of a concept/case file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for _ in range(5):
                line = f.readline().strip()
                if line.startswith("**Course:**"):
                    course_str = line.replace("**Course:**", "").strip()
                    return [c.strip() for c in course_str.split(",") if c.strip()]
    except Exception:
        pass
    return []


def load_existing_concepts(course_name=None, group_courses=None):
    """Load existing concept titles and slugs, tiered by course affinity."""
    if not WIKI_DIR.exists():
        return {} if course_name is None else {"same_course": {}, "same_group": {}, "other": {}}

    group_set = set(group_courses) if group_courses else set()
    same_course = {}
    same_group = {}
    other = {}

    for file in WIKI_DIR.glob(f"{CONCEPT_FILE_PREFIX}*{CONCEPT_FILE_SUFFIX}"):
        try:
            with open(file, 'r', encoding='utf-8') as f:
                first_line = f.readline().strip()
                if first_line.startswith("# "):
                    title = first_line[2:].strip()
                    slug = file.stem[len(CONCEPT_FILE_PREFIX):]

                    if course_name is None:
                        other[title] = slug
                        continue

                    file_courses = _extract_course_from_file(file)

                    if course_name in file_courses:
                        same_course[title] = slug
                    elif group_set & set(file_courses):
                        same_group[title] = slug
                    else:
                        other[title] = slug
        except Exception as e:
            print(f"   Error reading {file}: {e}")

    if course_name is None:
        return other

    return {"same_course": same_course, "same_group": same_group, "other": other}


def load_existing_cases(course_name=None, group_courses=None):
    """Load existing case titles and slugs, tiered by course affinity."""
    if not WIKI_DIR.exists():
        return {} if course_name is None else {"same_course": {}, "same_group": {}, "other": {}}

    group_set = set(group_courses) if group_courses else set()
    same_course = {}
    same_group = {}
    other = {}

    for file in WIKI_DIR.glob(f"{CASE_FILE_PREFIX}*{CONCEPT_FILE_SUFFIX}"):
        try:
            with open(file, 'r', encoding='utf-8') as f:
                first_line = f.readline().strip()
                if first_line.startswith("# Case: "):
                    title = first_line[len("# Case: "):].strip()
                elif first_line.startswith("# "):
                    title = first_line[2:].strip()
                else:
                    continue

                slug = file.stem[len(CASE_FILE_PREFIX):]

                if course_name is None:
                    other[title] = slug
                    continue

                file_courses = _extract_course_from_file(file)

                if course_name in file_courses:
                    same_course[title] = slug
                elif group_set & set(file_courses):
                    same_group[title] = slug
                else:
                    other[title] = slug
        except Exception as e:
            print(f"   Error reading {file}: {e}")

    if course_name is None:
        return other

    return {"same_course": same_course, "same_group": same_group, "other": other}


# ══════════════════════════════════════════════════════════════════════════════
#  Gemini API
# ══════════════════════════════════════════════════════════════════════════════

def setup_gemini():
    """Configure Gemini API. Returns True on success. Caches to avoid re-init."""
    global _gemini_configured
    if _gemini_configured:
        return True

    if not GEMINI_API_KEY:
        print("Gemini_Api_Key not found in .env")
        return False
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        print(f"   Gemini API configured (model: {MODEL})")
        _gemini_configured = True
        return True
    except Exception as e:
        print(f"   Gemini setup failed: {e}")
        return False


def call_gemini(prompt, timeout=180):
    """Call Gemini API. Returns response text or None."""
    import google.api_core.exceptions

    total_api_calls = 0

    for i, model_name in enumerate(GEMINI_MODELS):
        total_api_calls += 1
        try:
            model = genai.GenerativeModel(model_name)

            generation_config = {
                'temperature': 0.7,
                'top_p': 0.95,
                'top_k': 40,
                'max_output_tokens': 65536,
            }

            request_options = {
                'timeout': timeout,
                'retry': None
            }

            print(f"   Calling {model_name} (timeout: {timeout}s, retries: DISABLED)...")
            response = model.generate_content(
                prompt,
                generation_config=generation_config,
                request_options=request_options
            )

            if i == 0:
                print(f"   Success in {model_name} (1 API call)")
            else:
                print(f"   Success in fallback model: {model_name} ({total_api_calls} API calls total)")
            return response.text

        except google.api_core.exceptions.DeadlineExceeded as e:
            print(f"\n   TIMEOUT ERROR on {model_name}")
            print(f"   Error: {str(e)}")
            print(f"   Request exceeded {timeout}s timeout")
            print(f"   API calls made: {total_api_calls}")
            raise Exception(f"Request timed out after {timeout}s ({total_api_calls} API call consumed)")

        except google.api_core.exceptions.ResourceExhausted as e:
            print(f"\n   RATE LIMIT ERROR on {model_name}")
            print(f"   Error: {str(e)}")
            print(f"   API calls made so far: {total_api_calls}")
            if i < len(GEMINI_MODELS) - 1:
                print(f"   Trying fallback model: {GEMINI_MODELS[i+1]} (different quota pool)\n")
                continue
            else:
                print(f"   All {len(GEMINI_MODELS)} models exhausted ({total_api_calls} API calls total)")
                raise

        except google.api_core.exceptions.ServiceUnavailable as e:
            print(f"\n   SERVICE UNAVAILABLE (503) on {model_name}")
            print(f"   Error: {str(e)[:100]}")
            print(f"   API calls made so far: {total_api_calls}")
            if i < len(GEMINI_MODELS) - 1:
                print(f"   Trying fallback model: {GEMINI_MODELS[i+1]}\n")
                continue
            else:
                print(f"   All {len(GEMINI_MODELS)} models unavailable ({total_api_calls} API calls total)")
                raise

        except google.api_core.exceptions.InvalidArgument as e:
            print(f"\n   INVALID ARGUMENT ERROR on {model_name}")
            print(f"   Error: {str(e)}")
            print(f"   API calls made: {total_api_calls} (stopping - affects all models)")
            raise

        except google.api_core.exceptions.PermissionDenied as e:
            print(f"\n   PERMISSION DENIED ERROR")
            print(f"   Error: {str(e)}")
            raise

        except google.api_core.exceptions.Unauthenticated as e:
            print(f"\n   AUTHENTICATION ERROR")
            print(f"   Error: {str(e)}")
            raise

        except Exception as e:
            error_type = type(e).__name__
            error_msg = str(e)
            print(f"\n   UNEXPECTED ERROR on {model_name}")
            print(f"   Type: {error_type}")
            print(f"   Error: {error_msg}")
            print(f"   API calls made: {total_api_calls}")

            error_str = error_msg.lower()
            is_network_error = any(term in error_str for term in [
                'connection', 'network', 'unreachable', 'dns', 'socket', 'timeout'
            ])

            if is_network_error:
                print(f"   Network/connection issue (stopping - affects all models)")
                raise

            print(f"   Failing immediately to avoid wasting API calls")
            raise

    return None


# ══════════════════════════════════════════════════════════════════════════════
#  Text & Image Extraction
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_file(file_path):
    """Extract text from PDF, DOCX, or TXT file."""
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    elif suffix == '.pdf':
        if not pymupdf:
            print("pymupdf not installed. Install with: pip install pymupdf")
            return None

        try:
            text = ""
            doc = pymupdf.open(file_path)
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    elif suffix == '.docx':
        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return None

        try:
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None

    elif suffix == '.doc':
        raise ValueError(
            f"Legacy .doc files are not supported. Convert to .docx or .pdf first: {file_path.name}"
        )

    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")


def extract_images_from_pdf(file_path, output_dir):
    """Extract images/charts from PDF and save as PNG files."""
    if not pymupdf:
        print("   pymupdf not available - skipping image extraction")
        return []

    file_path = Path(file_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    extracted_images = []

    try:
        doc = pymupdf.open(file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    if len(image_bytes) < 10000:
                        continue

                    pdf_name = file_path.stem
                    image_filename = f"{pdf_name}_Page{page_num+1}_Plot{img_index}.png"
                    image_filepath = output_dir / image_filename

                    with open(image_filepath, "wb") as f:
                        f.write(image_bytes)

                    extracted_images.append(image_filename)
                    print(f"   Saved: {image_filename}")
                except Exception as e:
                    print(f"   Could not extract image on page {page_num+1}: {e}")

        return extracted_images

    except Exception as e:
        print(f"   Error extracting images: {e}")
        return []


# ══════════════════════════════════════════════════════════════════════════════
#  Helpers: slug, tiers, scoping
# ══════════════════════════════════════════════════════════════════════════════

def make_slug(title):
    """Create a URL-friendly slug from a title."""
    return (title.lower()
            .replace(" ", "-")
            .replace("(", "")
            .replace(")", "")
            .replace("/", "-")
            .replace("\u2014", "-")
            .replace("'", "")
            .replace(":", "")
            .replace(",", ""))


def flatten_tiers(tiered):
    """Flatten a tiered dict into a single {title: slug} dict."""
    if "same_course" not in tiered:
        return tiered
    flat = {}
    flat.update(tiered["other"])
    flat.update(tiered["same_group"])
    flat.update(tiered["same_course"])
    return flat


def scoped_concepts_for_prompt(tiered):
    """Return the concept dict to embed in LLM prompts (same_course + same_group only)."""
    if "same_course" not in tiered:
        return tiered
    scoped = {}
    scoped.update(tiered["same_group"])
    scoped.update(tiered["same_course"])
    return scoped


def same_course_only(tiered):
    """Return just the same_course tier. Falls back to flat dict if not tiered."""
    if "same_course" not in tiered:
        return tiered
    return tiered["same_course"]


def _fuzzy_match(concept_title, concepts_dict):
    """Fuzzy match a title against a dict of {title: slug}. Returns slug or None."""
    lower_title = concept_title.lower()
    for existing_title, slug in concepts_dict.items():
        existing_lower = existing_title.lower()
        if (lower_title in existing_lower or
            existing_lower in lower_title or
            len(set(lower_title.split()) & set(existing_lower.split())) >= 2):
            return slug
    return None


def check_for_duplicates(concept_title, existing_concepts):
    """Check if concept already exists using tiered matching."""
    # Legacy flat dict mode
    if "same_course" not in existing_concepts:
        if concept_title in existing_concepts:
            return existing_concepts[concept_title]
        return _fuzzy_match(concept_title, existing_concepts)

    same_course = existing_concepts["same_course"]
    same_group = existing_concepts["same_group"]
    other = existing_concepts["other"]

    # 1. Same course: exact + fuzzy
    if concept_title in same_course:
        return same_course[concept_title]
    fuzzy = _fuzzy_match(concept_title, same_course)
    if fuzzy:
        return fuzzy

    # 2. Same group: exact only
    if concept_title in same_group:
        return same_group[concept_title]

    # 3. Other: exact only
    if concept_title in other:
        return other[concept_title]

    return None


def check_case_duplicates(case_name, existing_cases):
    """Check if case already exists (exact or similar)."""
    flat = flatten_tiers(existing_cases)

    if case_name in flat:
        return flat[case_name]

    stop_words = {'a', 'an', 'the', 'of', 'in', 'on', 'at', 'to', 'for', 'and', 'or', 'is', 'are', 'was', 'were', 'by', 'from', 'with', 'as', 'its', 'it', 'that', 'this', 'not', 'but', 'be', 'has', 'had', 'have', 'do', 'does', 'did', 'will', 'can', 'may', 'about', 'into', 'over', 'after', 'before', 'between', 'under', 'model', 'case', 'study'}
    lower_name = case_name.lower()
    for existing_name, slug in flat.items():
        existing_lower = existing_name.lower()
        if (lower_name in existing_lower or existing_lower in lower_name):
            return slug
        new_words = {w for w in lower_name.split() if len(w) >= 4 and w not in stop_words}
        old_words = {w for w in existing_lower.split() if len(w) >= 4 and w not in stop_words}
        if len(new_words & old_words) >= 2:
            return slug

    return None


# ══════════════════════════════════════════════════════════════════════════════
#  Seed Concepts
# ══════════════════════════════════════════════════════════════════════════════

def seed_concepts(course_name):
    """Create stub concept files from seed_concepts in courses.json. No API calls."""
    if not COURSES_FILE.exists():
        print("   Error: courses.json not found")
        return 0

    with open(COURSES_FILE, 'r', encoding='utf-8') as f:
        courses = json.load(f)

    if course_name not in courses:
        print(f"   Error: Course '{course_name}' not found in courses.json")
        return 0

    seeds = courses[course_name].get("seed_concepts", [])
    if not seeds:
        print(f"   No seed_concepts defined for '{course_name}' in courses.json")
        return 0

    print(f"\n   Seeding {len(seeds)} concepts for {course_name}...")
    created = 0
    skipped = 0

    for title in seeds:
        slug = make_slug(title)
        filepath = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

        if filepath.exists():
            print(f"   SKIP: {title} (already exists)")
            skipped += 1
            continue

        course_tag = course_name.lower().replace(" ", "")
        markdown = f"""# {title}

**Course:** {course_name}
**Source:** Seed concept

## Definition

(To be populated from lecture material)

## Key Points

(To be populated from lecture material)

## Formulas & Equations

(No formulas)

## Examples

(To be populated from lecture material)

## Related Concepts

(No related concepts in wiki yet)

## Notes

Seed concept — will be automatically enriched when lecture material is processed.

## References

- {course_name}: Seed concept

**Tags:** #concept #{course_tag}
**Status:** Seed
"""
        WIKI_DIR.mkdir(parents=True, exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(markdown)
        print(f"   CREATED: Concept-{slug}.md")
        created += 1

    print(f"\n   Seed complete: {created} created, {skipped} skipped (already exist)")
    return created


# ══════════════════════════════════════════════════════════════════════════════
#  LLM Extraction & Markdown Generation
# ══════════════════════════════════════════════════════════════════════════════

def extract_concepts_with_llm(file_content, file_name, existing_concepts, timeout=180):
    """Single Gemini API call to extract ALL concepts from file content."""
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"

    prompt = f"""You are extracting economics lecture content for a wiki.

FILE: {file_name}

EXISTING WIKI CONCEPTS (use these exact names for [[Wikilinks]]):
{existing_list}

TASK:
Extract ALL distinct concepts taught in this lecture. A typical lecture covers 20-25 concepts.
For EACH concept, provide complete details.

FULL LECTURE CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "concepts": [
    {{
      "primary_concept": "Concept Title",
      "definition": "Clear definition in 2-3 sentences",
      "key_points": ["point 1", "point 2", "point 3"],
      "examples": ["example 1", "example 2"],
      "formulas": ["formula 1", "formula 2"],
      "existing_wikilinks": ["Existing Concept 1", "Existing Concept 2"],
      "related_concepts": ["Related Concept 1", "Related Concept 2"],
      "source": "{file_name}",
      "notes": "Any important context"
    }}
  ]
}}

IMPORTANT:
- Extract EVERY concept, not just the main one
- Each concept should be specific (e.g. "Price Elasticity of Demand" not just "Elasticity")
- Include definitions, formulas, and examples for each
- Reference existing wiki concepts using their exact names for [[Wikilinks]]
"""

    try:
        response_text = call_gemini(prompt, timeout=timeout)
        if not response_text:
            print(f"   All models failed")
            return None

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None

        json_str = response_text[start_idx:end_idx]
        result = repair_json(json_str)
        if result is None:
            print(f"   JSON parsing error (repair failed)")
            print(f"Response text: {response_text[:500]}")
            return None
        return result.get("concepts", [])
    except Exception as e:
        print(f"   LLM error: {e}")
        return None


def create_concept_markdown(concept_data, existing_concepts, course_name=None):
    """Create markdown file with proper [[Wikilinks]]."""
    title = concept_data["primary_concept"]
    slug = make_slug(title)

    flat_concepts = flatten_tiers(existing_concepts)

    wikilinks = []
    for concept in concept_data.get("existing_wikilinks", []):
        if concept in flat_concepts:
            wikilinks.append(f"- [[{concept}]]")

    for concept in concept_data.get("related_concepts", []):
        if concept in flat_concepts and concept not in concept_data.get("existing_wikilinks", []):
            wikilinks.append(f"- [[{concept}]]")

    wikilinks_section = "\n".join(wikilinks) if wikilinks else "(No related concepts in wiki yet)"

    formulas = concept_data.get("formulas", [])
    formulas_section = "\n".join(f"- {f}" for f in formulas) if formulas else "(No formulas)"

    examples = concept_data.get("examples", [])
    examples_section = "\n".join(f"- {e}" for e in examples) if examples else "(No examples)"

    key_points = concept_data.get("key_points", [])
    key_points_section = "\n".join(f"- {p}" for p in key_points) if key_points else "(No key points)"

    source = concept_data.get("source", "Lecture material")
    course_line = f"\n**Course:** {course_name}" if course_name else ""

    if course_name:
        reference = f"- {course_name}: {source}"
    else:
        reference = f"- {source}"

    markdown = f"""# {title}
{course_line}
**Source:** {source}

## Definition

{concept_data.get("definition", "(No definition extracted)")}

## Key Points

{key_points_section}

## Formulas & Equations

{formulas_section}

## Examples

{examples_section}

## Related Concepts

{wikilinks_section}

## Notes

{concept_data.get("notes", "Extracted from lecture material.")}

## References

{reference}
"""

    return slug, markdown


def merge_concepts_with_llm(merge_requests, existing_concepts, course_name=None, timeout=180):
    """Batch merge: send all duplicate concepts to Gemini in 1 API call."""
    if not merge_requests:
        return {}

    merge_items = []
    for slug, data in merge_requests.items():
        merge_items.append(f"""
--- CONCEPT TO MERGE: {slug} ---
EXISTING MARKDOWN:
{data['existing_markdown']}

NEW CONTENT TO ADD:
Title: {data['new_data']['primary_concept']}
Definition: {data['new_data'].get('definition', '')}
Key Points: {json.dumps(data['new_data'].get('key_points', []))}
Examples: {json.dumps(data['new_data'].get('examples', []))}
Formulas: {json.dumps(data['new_data'].get('formulas', []))}
Source: {data['new_data'].get('source', '')}
Notes: {data['new_data'].get('notes', '')}
Course: {course_name or 'Unknown'}
--- END {slug} ---
""")

    flat_concepts = flatten_tiers(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in flat_concepts.keys()) if flat_concepts else "(none)"

    course_merge_instruction = ""
    if course_name:
        course_merge_instruction = f"""- The **Course:** metadata line should list ALL courses this concept appears in, comma-separated (e.g. "**Course:** Microeconomics, Marketing")
- If the existing page already has a **Course:** line, append "{course_name}" if not already listed
- If no **Course:** line exists, add "**Course:** {course_name}" after the title
- References should be prefixed with course name (e.g. "- {course_name}: filename.pdf")"""

    prompt = f"""You are merging new lecture content into existing wiki concept pages.

For each concept below, rewrite the FULL markdown page that seamlessly integrates the new content.

RULES:
- Merge new key points, examples, formulas into the existing sections (don't create "Additional" subsections)
- Remove duplicates - if a point already exists, don't add it again
- Keep the same markdown structure: # Title, **Course:**, **Source:**, ## Definition, ## Key Points, ## Formulas & Equations, ## Examples, ## Related Concepts, ## Notes, ## References
- Add ALL sources to a ## References section at the bottom (list each source file)
- Use [[Wikilinks]] only for these existing concepts: {existing_list}
- Keep content comprehensive but not repetitive
{course_merge_instruction}

CONCEPTS TO MERGE:
{"".join(merge_items)}

RESPOND ONLY WITH JSON (no markdown wrapping):
{{
  "merged": {{
    "{list(merge_requests.keys())[0]}": "full rewritten markdown here...",
    ...
  }}
}}
"""

    try:
        response_text = call_gemini(prompt, timeout=timeout)
        if not response_text:
            print(f"   All models failed for merge")
            return {}

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        result = repair_json(response_text[start_idx:end_idx])
        if result is None:
            print(f"   Merge JSON repair failed")
            return {}
        return result.get('merged', {})

    except Exception as e:
        print(f"   Merge API error: {e}")
        return {}


# ══════════════════════════════════════════════════════════════════════════════
#  Case Study Extraction
# ══════════════════════════════════════════════════════════════════════════════

def extract_case_with_llm(file_content, file_name, existing_concepts, timeout=180):
    """Single Gemini API call to extract case study data."""
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"

    prompt = f"""You are extracting a business case study for a wiki.

FILE: {file_name}

EXISTING WIKI CONCEPTS (reference these using [[Wikilinks]]):
{existing_list}

TASK:
Read the provided business case study or simulation brief. Extract the critical data required for a strategic debrief.

FULL CASE CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "case_name": "The formal case name (e.g. 'Heidi Roizen at KPCB')",
  "core_dilemma": "Two to three sentences defining the primary strategic, leadership, or operational problem the protagonist is facing.",
  "stakeholders": [
    {{
      "name": "Person/Entity Name",
      "role": "Their role in the case",
      "incentives": "Their primary motivations or pressure points"
    }}
  ],
  "financial_context": "Any relevant margins, valuation multiples, budgets, or macro-economic constraints mentioned. If none, say 'No specific financial data provided.'",
  "related_concepts": ["Existing Concept 1", "Existing Concept 2"],
  "course": "",
  "source": "{file_name}"
}}

IMPORTANT:
- The case_name should be the formal title of the case, not a generic description
- Stakeholders should include ALL key actors, their roles, and motivations
- Financial context should include specific numbers when available
- related_concepts should ONLY reference concepts from the existing wiki list above
"""

    try:
        response_text = call_gemini(prompt, timeout=timeout)
        if not response_text:
            print(f"   All models failed")
            return None

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None

        json_str = response_text[start_idx:end_idx]
        result = repair_json(json_str)
        if result is None:
            print(f"   JSON parsing error (repair failed)")
            return None
        return result
    except Exception as e:
        print(f"   LLM error: {e}")
        return None


def create_case_markdown(case_data, existing_concepts, course_name=None):
    """Create Case-*.md file with schema.md sections."""
    case_name = case_data["case_name"]
    slug = make_slug(case_name)

    flat_concepts = flatten_tiers(existing_concepts)

    stakeholders = case_data.get("stakeholders", [])
    if stakeholders:
        stakeholders_section = "\n".join(
            f"- **{s['name']}** ({s.get('role', 'Unknown role')}): {s.get('incentives', '')}"
            for s in stakeholders
        )
    else:
        stakeholders_section = "(No stakeholders identified)"

    wikilinks = []
    for concept in case_data.get("related_concepts", []):
        if concept in flat_concepts:
            wikilinks.append(f"- [[{concept}]]")
    wikilinks_section = "\n".join(wikilinks) if wikilinks else "(No related concepts in wiki yet)"

    source = case_data.get("source", "Case material")
    course_line = f"\n**Course:** {course_name}" if course_name else ""

    if course_name:
        reference = f"- {course_name}: {source}"
    else:
        reference = f"- {source}"

    markdown = f"""# Case: {case_name}
**Tags:** #case-study #unresolved
{course_line}
**Source:** {source}

## 1. Core Dilemma

{case_data.get("core_dilemma", "(No dilemma extracted)")}

## 2. Key Stakeholders & Incentives

{stakeholders_section}

## 3. Financial Context & Constraints

{case_data.get("financial_context", "(No financial context extracted)")}

## 4. Class Discussion & Takeaways

*(To be populated by transcript processing)*

## Related Concepts

{wikilinks_section}

## References

{reference}
"""

    return slug, markdown


# ══════════════════════════════════════════════════════════════════════════════
#  Transcript Extraction & Case Discussion Updates
# ══════════════════════════════════════════════════════════════════════════════

def extract_transcript_with_llm(file_content, file_name, existing_concepts, existing_cases, timeout=180):
    """Single Gemini API call to extract BOTH concepts and case discussions from a transcript."""
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    prompt_cases = same_course_only(existing_cases)
    existing_concept_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"
    existing_case_list = "\n".join(f"- {title}" for title in prompt_cases.keys()) if prompt_cases else "(none yet)"

    prompt = f"""You are processing a class transcript for a wiki. This transcript may contain:
1. New concepts being taught (same as a lecture)
2. Discussion of case studies that were previously assigned

FILE: {file_name}

EXISTING WIKI CONCEPTS (use these exact names for [[Wikilinks]]):
{existing_concept_list}

EXISTING CASE STUDIES (match discussions to these case names):
{existing_case_list}

TASK:
Extract BOTH concepts AND case study discussions from this transcript.

FULL TRANSCRIPT CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "concepts": [
    {{
      "primary_concept": "Concept Title",
      "definition": "Clear definition in 2-3 sentences",
      "key_points": ["point 1", "point 2", "point 3"],
      "examples": ["example 1", "example 2"],
      "formulas": ["formula 1", "formula 2"],
      "existing_wikilinks": ["Existing Concept 1", "Existing Concept 2"],
      "related_concepts": ["Related Concept 1", "Related Concept 2"],
      "source": "{file_name}",
      "notes": "Any important context"
    }}
  ],
  "case_discussions": [
    {{
      "case_name": "Exact Case Name from existing cases list",
      "discussion_summary": "2-3 paragraph summary of how this case was discussed in class",
      "key_takeaways": ["takeaway 1", "takeaway 2"],
      "student_perspectives": ["perspective 1", "perspective 2"],
      "professor_insights": ["insight 1", "insight 2"]
    }}
  ]
}}

IMPORTANT:
- Extract ALL concepts taught, just like a lecture
- For case_discussions, ONLY include cases that were actually discussed in this transcript
- case_name must match an existing case from the list above (use the exact name)
- If no cases were discussed, return an empty case_discussions array
- student_perspectives: notable student arguments, debates, or questions
- professor_insights: key points the professor emphasized or corrected
"""

    try:
        response_text = call_gemini(prompt, timeout=timeout)
        if not response_text:
            print(f"   All models failed")
            return None, None

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None, None

        json_str = response_text[start_idx:end_idx]
        result = repair_json(json_str)
        if result is None:
            print(f"   JSON parsing error (repair failed)")
            return None, None
        return result.get("concepts", []), result.get("case_discussions", [])
    except Exception as e:
        print(f"   LLM error: {e}")
        return None, None


def update_case_discussion(case_slug, discussion_data, source_file):
    """Update a Case-*.md file's '## 4. Class Discussion & Takeaways' section."""
    case_file = WIKI_DIR / f"{CASE_FILE_PREFIX}{case_slug}{CONCEPT_FILE_SUFFIX}"

    if not case_file.exists():
        print(f"   Case file not found: {case_file}")
        return False

    try:
        with open(case_file, 'r', encoding='utf-8') as f:
            content = f.read()

        discussion_md = discussion_data.get("discussion_summary", "")

        takeaways = discussion_data.get("key_takeaways", [])
        if takeaways:
            discussion_md += "\n\n### Key Takeaways\n\n"
            discussion_md += "\n".join(f"- {t}" for t in takeaways)

        student_perspectives = discussion_data.get("student_perspectives", [])
        if student_perspectives:
            discussion_md += "\n\n### Student Perspectives\n\n"
            discussion_md += "\n".join(f"- {p}" for p in student_perspectives)

        professor_insights = discussion_data.get("professor_insights", [])
        if professor_insights:
            discussion_md += "\n\n### Professor Insights\n\n"
            discussion_md += "\n".join(f"- {i}" for i in professor_insights)

        discussion_md += f"\n\n*Source: {source_file}*"

        placeholder = "*(To be populated by transcript processing)*"
        section_header = "## 4. Class Discussion & Takeaways"

        if placeholder in content:
            content = content.replace(placeholder, discussion_md)
        elif section_header in content:
            section_start = content.index(section_header) + len(section_header)
            next_section = content.find("\n## ", section_start)

            if next_section != -1:
                existing_content = content[section_start:next_section].rstrip()
                content = (content[:section_start] +
                          existing_content + "\n\n---\n\n" + discussion_md + "\n\n" +
                          content[next_section:])
            else:
                content = content.rstrip() + "\n\n---\n\n" + discussion_md + "\n"

        content = content.replace("#unresolved", "#resolved")

        with open(case_file, 'w', encoding='utf-8') as f:
            f.write(content)

        return True

    except Exception as e:
        print(f"   Error updating case discussion: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  Main Processing Flows (lecture, case, transcript)
# ══════════════════════════════════════════════════════════════════════════════

class ProcessingError(Exception):
    """Raised when file processing fails (replaces sys.exit(1) from process_single_file)."""
    pass


def process_lecture_file(content, file_path, existing_concepts, course_name, timeout=180):
    """Process a lecture file: extract concepts, create/merge."""
    api_calls = 0

    print(f"\n   Using Gemini to extract concepts (1 API call)...")
    concepts_list = extract_concepts_with_llm(content, Path(file_path).name, existing_concepts, timeout=timeout)
    api_calls += 1

    if not concepts_list:
        raise ProcessingError("Failed to extract concept data")

    print(f"   Extracted {len(concepts_list)} concepts")

    print(f"\n{'='*70}")
    print(f"   Extraction Results:")
    print(f"{'='*70}")
    for i, concept_data in enumerate(concepts_list, 1):
        print(f"\n  {i}. {concept_data['primary_concept']}")
        print(f"     Definition: {concept_data.get('definition', '')[:80]}...")
        print(f"     Key Points: {len(concept_data.get('key_points', []))} | Examples: {len(concept_data.get('examples', []))} | Formulas: {len(concept_data.get('formulas', []))}")

    created = 0
    updated = 0
    merge_requests = {}

    for concept_data in concepts_list:
        title = concept_data["primary_concept"]
        print(f"\n{'~'*50}")
        print(f"   Processing: {title}")

        duplicate = check_for_duplicates(title, existing_concepts)

        if duplicate:
            print(f"   Duplicate found -> queuing for merge (Concept-{duplicate}.md)")
            existing_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{duplicate}{CONCEPT_FILE_SUFFIX}"

            try:
                with open(existing_file, 'r', encoding='utf-8') as f:
                    existing_markdown = f.read()

                if duplicate in merge_requests:
                    prev = merge_requests[duplicate]['new_data']
                    prev['key_points'] = prev.get('key_points', []) + concept_data.get('key_points', [])
                    prev['examples'] = prev.get('examples', []) + concept_data.get('examples', [])
                    prev['formulas'] = prev.get('formulas', []) + concept_data.get('formulas', [])
                else:
                    merge_requests[duplicate] = {
                        'existing_markdown': existing_markdown,
                        'new_data': concept_data
                    }
            except Exception as e:
                print(f"   Error reading existing file: {e}")
        else:
            slug, markdown = create_concept_markdown(concept_data, existing_concepts, course_name)
            output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(markdown)
                print(f"   Created: {output_file}")
                created += 1
                if "same_course" in existing_concepts:
                    existing_concepts["same_course"][title] = slug
                else:
                    existing_concepts[title] = slug
            except Exception as e:
                print(f"   Error saving: {e}")

    if merge_requests:
        print(f"\n{'~'*50}")
        print(f"   Merging {len(merge_requests)} existing concepts with new content (1 API call)...")
        api_calls += 1

        merged_results = merge_concepts_with_llm(merge_requests, existing_concepts, course_name, timeout=timeout)

        for slug, new_markdown in merged_results.items():
            output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(new_markdown)
                print(f"   Merged: {output_file}")
                updated += 1
            except Exception as e:
                print(f"   Error saving merge: {e}")

    print(f"\n   Created: {created} new concepts")
    print(f"   Updated: {updated} existing concepts")
    return api_calls


def process_case_file(content, file_path, existing_concepts, existing_cases, course_name, timeout=180):
    """Process a case study file: extract case data, create Case-*.md."""
    api_calls = 0

    print(f"\n   Using Gemini to extract case study (1 API call)...")
    case_data = extract_case_with_llm(content, Path(file_path).name, existing_concepts, timeout=timeout)
    api_calls += 1

    if not case_data:
        raise ProcessingError("Failed to extract case data")

    case_name = case_data.get("case_name", "Unknown Case")
    print(f"   Extracted case: {case_name}")

    print(f"\n{'='*70}")
    print(f"   Case Study Results:")
    print(f"{'='*70}")
    print(f"   Name: {case_name}")
    print(f"   Dilemma: {case_data.get('core_dilemma', '')[:100]}...")
    print(f"   Stakeholders: {len(case_data.get('stakeholders', []))}")
    print(f"   Related concepts: {len(case_data.get('related_concepts', []))}")

    slug, markdown = create_case_markdown(case_data, existing_concepts, course_name)
    output_file = WIKI_DIR / f"{CASE_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown)
        print(f"\n   Created: {output_file}")
        if "same_course" in existing_cases:
            existing_cases["same_course"][case_name] = slug
        else:
            existing_cases[case_name] = slug
    except Exception as e:
        print(f"   Error saving case: {e}")

    return api_calls


def process_transcript_file(content, file_path, existing_concepts, existing_cases, course_name, timeout=180):
    """Process a transcript: extract concepts AND update case discussions."""
    api_calls = 0

    print(f"\n   Using Gemini to extract concepts + case discussions (1 API call)...")
    concepts_list, case_discussions = extract_transcript_with_llm(
        content, Path(file_path).name, existing_concepts, existing_cases, timeout=timeout
    )
    api_calls += 1

    if concepts_list is None:
        raise ProcessingError("Failed to extract transcript data")

    print(f"   Extracted {len(concepts_list)} concepts, {len(case_discussions or [])} case discussions")

    # Part 1: Process concepts (reuse lecture logic)
    if concepts_list:
        print(f"\n{'='*70}")
        print(f"   Concept Extraction Results:")
        print(f"{'='*70}")
        for i, concept_data in enumerate(concepts_list, 1):
            print(f"\n  {i}. {concept_data['primary_concept']}")
            print(f"     Definition: {concept_data.get('definition', '')[:80]}...")

        created = 0
        updated = 0
        merge_requests = {}

        for concept_data in concepts_list:
            title = concept_data["primary_concept"]
            print(f"\n{'~'*50}")
            print(f"   Processing: {title}")

            duplicate = check_for_duplicates(title, existing_concepts)

            if duplicate:
                print(f"   Duplicate found -> queuing for merge (Concept-{duplicate}.md)")
                existing_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{duplicate}{CONCEPT_FILE_SUFFIX}"

                try:
                    with open(existing_file, 'r', encoding='utf-8') as f:
                        existing_markdown = f.read()

                    if duplicate in merge_requests:
                        prev = merge_requests[duplicate]['new_data']
                        prev['key_points'] = prev.get('key_points', []) + concept_data.get('key_points', [])
                        prev['examples'] = prev.get('examples', []) + concept_data.get('examples', [])
                        prev['formulas'] = prev.get('formulas', []) + concept_data.get('formulas', [])
                    else:
                        merge_requests[duplicate] = {
                            'existing_markdown': existing_markdown,
                            'new_data': concept_data
                        }
                except Exception as e:
                    print(f"   Error reading existing file: {e}")
            else:
                slug, markdown = create_concept_markdown(concept_data, existing_concepts, course_name)
                output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

                try:
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(markdown)
                    print(f"   Created: {output_file}")
                    created += 1
                    if "same_course" in existing_concepts:
                        existing_concepts["same_course"][title] = slug
                    else:
                        existing_concepts[title] = slug
                except Exception as e:
                    print(f"   Error saving: {e}")

        if merge_requests:
            print(f"\n{'~'*50}")
            print(f"   Merging {len(merge_requests)} existing concepts with new content (1 API call)...")
            api_calls += 1

            merged_results = merge_concepts_with_llm(merge_requests, existing_concepts, course_name, timeout=timeout)

            for slug, new_markdown in merged_results.items():
                output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"
                try:
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(new_markdown)
                    print(f"   Merged: {output_file}")
                    updated += 1
                except Exception as e:
                    print(f"   Error saving merge: {e}")

        print(f"\n   Concepts — Created: {created}, Updated: {updated}")

    # Part 2: Update case discussions (file I/O only, no API calls)
    if case_discussions:
        print(f"\n{'='*70}")
        print(f"   Case Discussion Updates:")
        print(f"{'='*70}")

        discussions_updated = 0
        for disc in case_discussions:
            case_name = disc.get("case_name", "")
            print(f"\n   Updating case: {case_name}")

            case_slug = flatten_tiers(existing_cases).get(case_name)
            if not case_slug:
                case_slug = check_case_duplicates(case_name, existing_cases)

            if case_slug:
                success = update_case_discussion(case_slug, disc, Path(file_path).name)
                if success:
                    print(f"   Updated: Case-{case_slug}.md")
                    discussions_updated += 1
                else:
                    print(f"   Failed to update Case-{case_slug}.md")
            else:
                print(f"   No matching case found for: {case_name}")

        print(f"\n   Case discussions updated: {discussions_updated}")

    return api_calls


# ══════════════════════════════════════════════════════════════════════════════
#  Direct File Processing (replaces subprocess call)
# ══════════════════════════════════════════════════════════════════════════════

def process_file_directly(filepath, course_name, file_type="lectures", extract_images=False):
    """Process a single file directly (no subprocess).

    Args:
        filepath: Path to the file
        course_name: Course name
        file_type: "lectures", "cases", or "transcripts"
        extract_images: Whether to extract images from PDFs

    Returns:
        True on success, False on failure
    """
    # Map batch file_type to processor type
    type_map = {"lectures": "lecture", "cases": "case", "transcripts": "transcript"}
    proc_type = type_map.get(file_type, "lecture")

    print(f"\n{'='*70}")
    print(f"   Processing: {filepath}")
    print(f"   Course: {course_name}")
    print(f"   Type: {proc_type}")
    print(f"{'='*70}")

    try:
        # Setup Gemini (cached after first call)
        if not setup_gemini():
            return False

        # Load course groups and existing concepts/cases
        group_courses = load_course_groups(course_name)
        if group_courses:
            print(f"\n   Course group siblings: {', '.join(group_courses)}")

        print("\n   Loading existing concepts...")
        existing_concepts = load_existing_concepts(course_name, group_courses)
        flat_concepts = flatten_tiers(existing_concepts)
        sc = len(existing_concepts.get("same_course", {}))
        sg = len(existing_concepts.get("same_group", {}))
        ot = len(existing_concepts.get("other", {}))
        print(f"   Found {len(flat_concepts)} existing concepts (same_course={sc}, same_group={sg}, other={ot})")

        existing_cases = load_existing_cases(course_name, group_courses)
        flat_cases = flatten_tiers(existing_cases)
        sc_c = len(existing_cases.get("same_course", {}))
        sg_c = len(existing_cases.get("same_group", {}))
        ot_c = len(existing_cases.get("other", {}))
        print(f"   Found {len(flat_cases)} existing cases (same_course={sc_c}, same_group={sg_c}, other={ot_c})")

        # Extract images from PDF (optional)
        if proc_type in ('lecture', 'case') and Path(filepath).suffix.lower() == '.pdf' and extract_images:
            print(f"\n   Extracting charts from PDF (saving to assets folder)...")
            extracted_images = extract_images_from_pdf(filepath, WIKI_DIR / "assets" / "charts")
            if extracted_images:
                print(f"   Saved {len(extracted_images)} charts to MBAWiki/assets/charts/")
            else:
                print(f"   No charts found in PDF")
        elif not extract_images:
            print(f"\n   Image extraction: SKIPPED")

        # Extract text
        print(f"\n   Extracting text from file...")
        content = extract_text_from_file(filepath)
        if not content:
            print("   No content extracted")
            return False
        content_length = len(content)
        print(f"   Read {content_length:,} characters")

        # Smart timeout
        base_timeout = 180
        extra_timeout = (content_length // 100000) * 30
        smart_timeout = min(base_timeout + extra_timeout, 600)
        if smart_timeout > base_timeout:
            print(f"   Using extended timeout: {smart_timeout}s (content size: {content_length//1000}k chars)")

        # Dispatch based on type
        if proc_type == 'case':
            api_calls = process_case_file(content, filepath, existing_concepts, existing_cases, course_name, timeout=smart_timeout)
        elif proc_type == 'transcript':
            api_calls = process_transcript_file(content, filepath, existing_concepts, existing_cases, course_name, timeout=smart_timeout)
        else:
            api_calls = process_lecture_file(content, filepath, existing_concepts, course_name, timeout=smart_timeout)

        # Log the ingestion
        source_name = Path(filepath).name
        log_ingestion("ingest", source_name, proc_type, f"{api_calls} API calls")

        # Update search index incrementally
        try:
            from ingest.build_search_index import build_index  # noqa: E402
            total, updated = build_index(append_mode=True)
            print(f"   Search index: {total} entries ({updated} re-embedded) [OK]")
        except ImportError:
            print(f"   Search index: skipped (fastembed not installed)")
        except Exception as e:
            print(f"   Search index: failed (non-fatal): {e}")

        print(f"\n   Complete! API calls used: {api_calls}")
        return True

    except ProcessingError as e:
        print(f"   Processing failed: {e}")
        return False
    except (FileNotFoundError, ValueError) as e:
        print(f"   File error: {e}")
        return False
    except Exception as e:
        print(f"   Unexpected error: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  Google Drive Integration (from process_all_lite.py)
# ══════════════════════════════════════════════════════════════════════════════

def load_courses():
    with open(COURSES_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_tracker():
    if not TRACKER_FILE.exists():
        return {}
    with open(TRACKER_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_tracker(tracker):
    with open(TRACKER_FILE, 'w', encoding='utf-8') as f:
        json.dump(tracker, f, indent=2)


def is_file_processed(tracker, course_name, file_type, filename):
    return (course_name in tracker
            and file_type in tracker[course_name]
            and filename in tracker[course_name][file_type])


def mark_file_processed(tracker, course_name, file_type, filename):
    if course_name not in tracker:
        tracker[course_name] = {}
    if file_type not in tracker[course_name]:
        tracker[course_name][file_type] = {}
    tracker[course_name][file_type][filename] = datetime.now(timezone.utc).isoformat()
    save_tracker(tracker)


def setup_google_drive():
    from google.auth.transport.requests import Request
    from google_auth_oauthlib.flow import InstalledAppFlow

    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    creds = None

    token_path = str(PROJECT_ROOT / 'credentials' / 'token.json')
    creds_path = str(PROJECT_ROOT / 'credentials' / 'credentials.json')

    if os.path.exists(token_path):
        try:
            creds = Credentials.from_authorized_user_file(token_path, SCOPES)
        except Exception:
            creds = None

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                creds = None

        if not creds:
            try:
                flow = InstalledAppFlow.from_client_secrets_file(
                    creds_path, SCOPES)
                creds = flow.run_local_server(port=0)
            except Exception as e:
                print(f"Auth error: {e}")
                return None

        try:
            with open(token_path, 'w') as token:
                token.write(creds.to_json())
        except Exception:
            pass

    return build('drive', 'v3', credentials=creds)


def list_files_in_folder(service, folder_id):
    try:
        query = f"'{folder_id}' in parents and trashed=false"
        results = service.files().list(
            q=query, spaces='drive',
            fields='files(id, name, mimeType, modifiedTime)',
            pageSize=100, orderBy='modifiedTime desc'
        ).execute()
        return results.get('files', [])
    except Exception as e:
        print(f"Error listing files: {e}")
        return []


def get_week_number(filename):
    match = re.search(r'[Ww]eek\s*(\d+)', filename)
    if match:
        return int(match.group(1))
    match = re.search(r'(\d{1,2})_(\d{1,2})_(\d{4})', filename)
    if match:
        return int(match.group(1)) * 100 + int(match.group(2))
    return 999


def sort_files(files):
    pdfs = sorted([f for f in files if f['name'].lower().endswith('.pdf')],
                  key=lambda f: get_week_number(f['name']))
    docxs = sorted([f for f in files if f['name'].lower().endswith('.docx')],
                   key=lambda f: get_week_number(f['name']))
    txts = sorted([f for f in files if f['name'].lower().endswith('.txt')],
                  key=lambda f: get_week_number(f['name']))
    return pdfs + docxs + txts


def download_file(service, file_id, filename, course_dir):
    course_dir.mkdir(parents=True, exist_ok=True)
    output_file = course_dir / filename

    if output_file.exists():
        print(f"   Already downloaded: {filename}")
        return str(output_file)

    try:
        request = service.files().get_media(fileId=file_id)
        print(f"   Downloading: {filename}")
        with open(output_file, 'wb') as f:
            downloader = MediaIoBaseDownload(f, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
        print(f"   Downloaded: {output_file.stat().st_size / (1024*1024):.1f} MB")
        return str(output_file)
    except Exception as e:
        print(f"   Download failed: {e}")
        return None


def get_unprocessed(service, course_name, folder_id, file_type, tracker):
    """Get list of unprocessed files for a folder."""
    if not folder_id:
        return []
    files = list_files_in_folder(service, folder_id)
    sorted_files = sort_files(files)
    return [f for f in sorted_files
            if not is_file_processed(tracker, course_name, file_type, f['name'])]


# ══════════════════════════════════════════════════════════════════════════════
#  Batch Orchestration
# ══════════════════════════════════════════════════════════════════════════════

def run_batch(service, course_name, course_config, extract_images, wait_minutes):
    """Process all unprocessed files. Returns True if all done, False if error occurred."""
    lectures_folder_id = course_config['lectures_folder_id']
    cases_folder_id = course_config.get('cases_folder_id')
    transcripts_folder_id = course_config.get('transcripts_folder_id')
    course_dir = LOCAL_DIR / course_name

    # Seed concepts first (no API calls)
    seed_concepts(course_name)

    # Build the work queue: (folder_id, file_type, course_subdir)
    phases = [
        (lectures_folder_id, "lectures", course_dir),
        (cases_folder_id, "cases", course_dir / "cases"),
        (transcripts_folder_id, "transcripts", course_dir / "transcripts"),
    ]

    total_processed = 0

    for folder_id, file_type, subdir in phases:
        if not folder_id:
            continue

        tracker = load_tracker()
        unprocessed = get_unprocessed(service, course_name, folder_id, file_type, tracker)

        if not unprocessed:
            print(f"\n   [{file_type.upper()}] All done")
            continue

        print(f"\n{'='*60}")
        print(f"   [{file_type.upper()}] {len(unprocessed)} files to process")
        print(f"{'='*60}")

        for i, file_info in enumerate(unprocessed, 1):
            # Rate limit delay between files
            if total_processed > 0:
                delay = 20
                print(f"\n   Waiting {delay}s (rate limit)...")
                time.sleep(delay)

            print(f"\n   [{file_type.upper()} {i}/{len(unprocessed)}] {file_info['name']}")

            local_path = download_file(service, file_info['id'], file_info['name'], subdir)
            if not local_path:
                print(f"\n   Download failed. Will retry in {wait_minutes} minutes.")
                return False

            is_pdf = file_info['name'].lower().endswith('.pdf')
            use_images = extract_images and is_pdf

            success = process_file_directly(local_path, course_name,
                                            file_type=file_type, extract_images=use_images)

            if success:
                tracker = load_tracker()
                mark_file_processed(tracker, course_name, file_type, file_info['name'])
                total_processed += 1
                print(f"   Processed so far: {total_processed}")
            else:
                # Count remaining
                remaining_this_phase = len(unprocessed) - i
                remaining_later = 0
                for fid, ft, _ in phases[phases.index((folder_id, file_type, subdir))+1:]:
                    if fid:
                        remaining_later += len(get_unprocessed(
                            service, course_name, fid, ft, load_tracker()))
                total_remaining = remaining_this_phase + remaining_later

                print(f"\n{'='*60}")
                print(f"   ERROR on: {file_info['name']}")
                print(f"   Model: {MODEL} (no fallback)")
                print(f"   Processed so far: {total_processed}")
                print(f"   Remaining: {total_remaining}")
                print(f"   Will retry in {wait_minutes} minutes...")
                print(f"{'='*60}")
                return False

    return True


def main():
    courses = load_courses()

    # Parse args
    args = sys.argv[1:]
    course_name = None
    extract_images = '--images' in args
    wait_minutes = DEFAULT_WAIT_MINUTES

    if '--course' in args:
        idx = args.index('--course')
        if idx + 1 < len(args):
            course_name = args[idx + 1]

    if '--wait' in args:
        idx = args.index('--wait')
        if idx + 1 < len(args):
            wait_minutes = int(args[idx + 1])

    if not course_name:
        print("Usage: python process_standalone.py --course \"CourseName\"")
        print("       python process_standalone.py --course \"CourseName\" --images")
        print("       python process_standalone.py --course \"CourseName\" --wait 15")
        print(f"\nAvailable courses: {', '.join(courses.keys())}")
        sys.exit(1)

    if course_name not in courses:
        print(f"Course not found: {course_name}")
        print(f"Available: {', '.join(courses.keys())}")
        sys.exit(1)

    course_config = courses[course_name]

    print(f"{'='*60}")
    print(f"  STANDALONE BATCH PROCESSOR")
    print(f"  Course: {course_name}")
    print(f"  Model:  {MODEL} (no fallback)")
    print(f"  Retry:  {wait_minutes} min wait on error")
    print(f"  Images: {'ON' if extract_images else 'OFF'}")
    print(f"{'='*60}")

    service = setup_google_drive()
    if not service:
        sys.exit(1)

    attempt = 0
    while True:
        attempt += 1
        print(f"\n{'='*60}")
        print(f"  ATTEMPT #{attempt} — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"{'='*60}")

        all_done = run_batch(service, course_name, course_config, extract_images, wait_minutes)

        if all_done:
            print(f"\n{'='*60}")
            print(f"  ALL FILES PROCESSED")
            print(f"  Attempts: {attempt}")
            print(f"  Finished: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"{'='*60}")

            # Rebuild search index
            try:
                from ingest.build_search_index import build_index  # noqa: E402
                total, updated = build_index(append_mode=False)
                print(f"  Search index rebuilt: {total} entries")
            except ImportError:
                print(f"  Search index skipped (fastembed not installed)")
            except Exception as e:
                print(f"  Search index failed (non-fatal): {e}")

            print(f"\n  Next steps:")
            print(f"  1. python wiki_viewer/app.py")
            print(f"  2. http://127.0.0.1:5000/")
            break
        else:
            resume_time = datetime.now()
            resume_hour = (resume_time.hour + (resume_time.minute + wait_minutes) // 60) % 24
            resume_min = (resume_time.minute + wait_minutes) % 60
            print(f"\n  Sleeping {wait_minutes} minutes...")
            print(f"  Will resume at ~{resume_hour:02d}:{resume_min:02d}")
            print(f"  (Press Ctrl+C to stop)\n")

            try:
                time.sleep(wait_minutes * 60)
            except KeyboardInterrupt:
                print(f"\n  Stopped by user after {attempt} attempt(s).")
                sys.exit(0)


if __name__ == "__main__":
    main()
