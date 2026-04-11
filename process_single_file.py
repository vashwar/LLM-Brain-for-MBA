#!/usr/bin/env python3
"""
Process Single File for MBA Wiki
Processes ONE file at a time (PDF or TXT) and extracts ALL concepts.
Supports three file types: lectures, cases, and transcripts.
For PDFs: Extracts images to assets folder (for future use).
Uses a single Gemini API call per file to stay within rate limits.

Usage:
    python process_single_file.py "Transcript_class_lecture/week1.txt"
    python process_single_file.py "Transcript_class_lecture/slides.pdf" --course "Microeconomics"
    python process_single_file.py "case.pdf" --course "Leading People" --type case
    python process_single_file.py "transcript.txt" --course "Leading People" --type transcript
"""

import os
import sys
import json
import re
from pathlib import Path
from dotenv import load_dotenv
import google.generativeai as genai

try:
    import pymupdf
except ImportError:
    pymupdf = None

# Fix Unicode on Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# ── Configuration ────────────────────────────────────────────────────────────
load_dotenv()
GEMINI_API_KEY = os.getenv('Gemini_Api_Key')
WIKI_DIR = Path('MBAWiki')
CONCEPT_FILE_PREFIX = "Concept-"
CASE_FILE_PREFIX = "Case-"
CONCEPT_FILE_SUFFIX = ".md"
GEMINI_MODELS = [
    "gemini-3.1-flash-lite-preview",  # Highest free-tier rate limit
    "gemini-3-flash-preview",          # Fallback
    "gemini-2.5-flash",                # Final fallback
]
COURSE_GROUPS_FILE = Path('course_groups.json')
COURSES_FILE = Path('courses.json')
LOG_FILE = Path('log.md')


def log_ingestion(action, source_filename, file_type, details=""):
    """Append an entry to log.md with timestamp.

    Args:
        action: e.g., "ingest", "merge", "seed"
        source_filename: e.g., "Slides Week 1.pdf"
        file_type: e.g., "lecture", "case", "transcript"
        details: optional additional info
    """
    try:
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Format: ## [2026-04-10 17:30:00] {action} | {type}: {source} | {details}
        entry = f"## [{timestamp}] {action} | {file_type.capitalize()}: {source_filename}"
        if details:
            entry += f" | {details}"
        entry += "\n"

        # Append to log.md
        if LOG_FILE.exists():
            with open(LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(entry)
        else:
            # If log doesn't exist yet, create it with header
            header = "# Wiki Evolution Log\n\nAppend-only record of ingestions and updates.\n\n"
            with open(LOG_FILE, 'w', encoding='utf-8') as f:
                f.write(header)
                f.write(entry)
    except Exception as e:
        # Don't fail processing if logging fails
        print(f"   Warning: Could not log entry: {e}")


def load_course_groups(course_name):
    """Load course_groups.json and return sibling course names for the given course.
    Merges siblings from all matching groups (a course can be in multiple groups).
    Returns list of sibling course names (excluding self), or None if course not in any group.
    """
    if not COURSE_GROUPS_FILE.exists():
        return None

    try:
        with open(COURSE_GROUPS_FILE, 'r', encoding='utf-8') as f:
            groups = json.load(f)

        siblings = set()
        for group_name, courses in groups.items():
            if course_name in courses:
                siblings.update(c for c in courses if c != course_name)

        return sorted(siblings) if siblings else None
    except Exception as e:
        print(f"   Warning: Could not load course_groups.json: {e}")

    return None


def _extract_course_from_file(filepath):
    """Extract course names from the first 5 lines of a concept/case file.
    Returns list of course names, or empty list if none found.
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for _ in range(5):
                line = f.readline().strip()
                if line.startswith("**Course:**"):
                    course_str = line.replace("**Course:**", "").strip()
                    return [c.strip() for c in course_str.split(",") if c.strip()]
    except Exception:
        pass
    return []


def parse_args():
    """Parse command line arguments."""
    args = sys.argv[1:]
    parsed = {
        'file_path': None,
        'course': None,
        'no_images': '--no-images' in args,
        'type': 'lecture',
        'seed': '--seed' in args,
    }

    # Extract --course value
    if '--course' in args:
        idx = args.index('--course')
        if idx + 1 < len(args):
            parsed['course'] = args[idx + 1]

    # Extract --type value
    if '--type' in args:
        idx = args.index('--type')
        if idx + 1 < len(args):
            parsed['type'] = args[idx + 1]

    # Find file path (first positional arg that's not a flag or flag value)
    skip_next = False
    for i, arg in enumerate(args):
        if skip_next:
            skip_next = False
            continue
        if arg in ('--course', '--type'):
            skip_next = True
            continue
        if arg.startswith('--'):
            continue
        parsed['file_path'] = arg
        break

    # --seed doesn't require a file path
    if parsed['seed']:
        parsed['file_path'] = '__seed__'

    return parsed


def setup_gemini():
    """Configure Gemini API."""
    if not GEMINI_API_KEY:
        print("Gemini_Api_Key not found in .env")
        return False
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        print(f"   Gemini API configured (models: {' -> '.join(GEMINI_MODELS)})")
        return True
    except Exception as e:
        print(f"   Gemini setup failed: {e}")
        return False


def call_gemini(prompt):
    """Call Gemini API with model fallback chain. Returns response text or None."""
    for i, model_name in enumerate(GEMINI_MODELS):
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            if i == 0:
                pass  # Primary model worked, no extra logging
            else:
                print(f"   (using fallback model: {model_name})")
            return response.text
        except Exception as e:
            error_str = str(e).lower()
            is_rate_limit = ('429' in str(e) or 'rate' in error_str
                            or 'quota' in error_str or 'resource' in error_str)
            if is_rate_limit and i < len(GEMINI_MODELS) - 1:
                print(f"   Rate limited on {model_name}, falling back to {GEMINI_MODELS[i+1]}...")
                continue
            else:
                raise
    return None


def load_existing_concepts(course_name=None, group_courses=None):
    """Load existing concept titles and slugs, tiered by course affinity.

    When course_name is provided, returns:
        {"same_course": {title: slug}, "same_group": {title: slug}, "other": {title: slug}}

    When course_name is None (legacy mode), returns flat dict: {title: slug}
    """
    if not WIKI_DIR.exists():
        print(f"   Wiki directory not found: {WIKI_DIR}")
        return {} if course_name is None else {"same_course": {}, "same_group": {}, "other": {}}

    group_set = set(group_courses) if group_courses else set()
    same_course = {}
    same_group = {}
    other = {}

    for file in WIKI_DIR.glob(f"{CONCEPT_FILE_PREFIX}*{CONCEPT_FILE_SUFFIX}"):
        try:
            with open(file, 'r', encoding='utf-8') as f:
                first_line = f.readline().strip()
                if first_line.startswith("# "):
                    title = first_line[2:].strip()
                    slug = file.stem[len(CONCEPT_FILE_PREFIX):]

                    if course_name is None:
                        other[title] = slug
                        continue

                    file_courses = _extract_course_from_file(file)

                    if course_name in file_courses:
                        same_course[title] = slug
                    elif group_set & set(file_courses):
                        same_group[title] = slug
                    else:
                        other[title] = slug
        except Exception as e:
            print(f"   Error reading {file}: {e}")

    if course_name is None:
        return other

    return {"same_course": same_course, "same_group": same_group, "other": other}


def load_existing_cases(course_name=None, group_courses=None):
    """Load existing case titles and slugs, tiered by course affinity.

    When course_name is provided, returns:
        {"same_course": {title: slug}, "same_group": {title: slug}, "other": {title: slug}}

    When course_name is None (legacy mode), returns flat dict: {title: slug}
    """
    if not WIKI_DIR.exists():
        return {} if course_name is None else {"same_course": {}, "same_group": {}, "other": {}}

    group_set = set(group_courses) if group_courses else set()
    same_course = {}
    same_group = {}
    other = {}

    for file in WIKI_DIR.glob(f"{CASE_FILE_PREFIX}*{CONCEPT_FILE_SUFFIX}"):
        try:
            with open(file, 'r', encoding='utf-8') as f:
                first_line = f.readline().strip()
                if first_line.startswith("# Case: "):
                    title = first_line[len("# Case: "):].strip()
                elif first_line.startswith("# "):
                    title = first_line[2:].strip()
                else:
                    continue

                slug = file.stem[len(CASE_FILE_PREFIX):]

                if course_name is None:
                    other[title] = slug
                    continue

                file_courses = _extract_course_from_file(file)

                if course_name in file_courses:
                    same_course[title] = slug
                elif group_set & set(file_courses):
                    same_group[title] = slug
                else:
                    other[title] = slug
        except Exception as e:
            print(f"   Error reading {file}: {e}")

    if course_name is None:
        return other

    return {"same_course": same_course, "same_group": same_group, "other": other}


def extract_course_metadata(filepath):
    """Extract the Course: line from an existing concept file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        match = re.search(r'^\*\*Course:\*\*\s*(.+)$', content, re.MULTILINE)
        if match:
            return [c.strip() for c in match.group(1).split(',')]
    except Exception:
        pass
    return []


def extract_images_from_pdf(file_path, output_dir):
    """Extract images/charts from PDF and save as PNG files (for future use)."""
    if not pymupdf:
        print("   pymupdf not available - skipping image extraction")
        return []

    file_path = Path(file_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    extracted_images = []

    try:
        doc = pymupdf.open(file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Skip small images (logos, icons, decorative elements)
                    if len(image_bytes) < 10000:
                        continue

                    # Create clean filename
                    pdf_name = file_path.stem
                    image_filename = f"{pdf_name}_Page{page_num+1}_Plot{img_index}.png"
                    image_filepath = output_dir / image_filename

                    with open(image_filepath, "wb") as f:
                        f.write(image_bytes)

                    extracted_images.append(image_filename)
                    print(f"   Saved: {image_filename}")
                except Exception as e:
                    print(f"   Could not extract image on page {page_num+1}: {e}")

        return extracted_images

    except Exception as e:
        print(f"   Error extracting images: {e}")
        return []


def extract_text_from_file(file_path):
    """Extract text from PDF, DOCX, or TXT file."""
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    elif suffix == '.pdf':
        if not pymupdf:
            print("pymupdf not installed. Install with: pip install pymupdf")
            return None

        try:
            text = ""
            doc = pymupdf.open(file_path)
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    elif suffix == '.docx':
        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return None

        try:
            doc = Document(file_path)
            parts = []
            # Paragraphs
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            # Tables (flatten row cells tab-separated)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None

    elif suffix == '.doc':
        raise ValueError(
            f"Legacy .doc files are not supported. Convert to .docx or .pdf first: {file_path.name}"
        )

    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")


def make_slug(title):
    """Create a URL-friendly slug from a title."""
    return (title.lower()
            .replace(" ", "-")
            .replace("(", "")
            .replace(")", "")
            .replace("/", "-")
            .replace("—", "-")
            .replace("'", "")
            .replace(":", "")
            .replace(",", ""))


def seed_concepts(course_name):
    """Create stub concept files from seed_concepts in courses.json.
    Skips concepts that already have a file. No API calls needed.
    """
    if not COURSES_FILE.exists():
        print("   Error: courses.json not found")
        return 0

    with open(COURSES_FILE, 'r', encoding='utf-8') as f:
        courses = json.load(f)

    if course_name not in courses:
        print(f"   Error: Course '{course_name}' not found in courses.json")
        return 0

    seeds = courses[course_name].get("seed_concepts", [])
    if not seeds:
        print(f"   No seed_concepts defined for '{course_name}' in courses.json")
        return 0

    print(f"\n   Seeding {len(seeds)} concepts for {course_name}...")
    created = 0
    skipped = 0

    for title in seeds:
        slug = make_slug(title)
        filepath = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

        if filepath.exists():
            print(f"   SKIP: {title} (already exists)")
            skipped += 1
            continue

        course_tag = course_name.lower().replace(" ", "")
        markdown = f"""# {title}

**Course:** {course_name}
**Source:** Seed concept

## Definition

(To be populated from lecture material)

## Key Points

(To be populated from lecture material)

## Formulas & Equations

(No formulas)

## Examples

(To be populated from lecture material)

## Related Concepts

(No related concepts in wiki yet)

## Notes

Seed concept — will be automatically enriched when lecture material is processed.

## References

- {course_name}: Seed concept

**Tags:** #concept #{course_tag}
**Status:** Seed
"""
        WIKI_DIR.mkdir(parents=True, exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(markdown)
        print(f"   CREATED: Concept-{slug}.md")
        created += 1

    print(f"\n   Seed complete: {created} created, {skipped} skipped (already exist)")
    return created


# ── Lecture Processing ───────────────────────────────────────────────────────

def extract_concepts_with_llm(file_content, file_name, existing_concepts):
    """
    Single Gemini API call to extract ALL concepts from file content.
    Returns a list of concepts.
    """

    # Build list of existing concepts for the prompt (scoped to same course + group)
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"

    prompt = f"""You are extracting economics lecture content for a wiki.

FILE: {file_name}

EXISTING WIKI CONCEPTS (use these exact names for [[Wikilinks]]):
{existing_list}

TASK:
Extract ALL distinct concepts taught in this lecture. A typical lecture covers 20-25 concepts.
For EACH concept, provide complete details.

FULL LECTURE CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "concepts": [
    {{
      "primary_concept": "Concept Title",
      "definition": "Clear definition in 2-3 sentences",
      "key_points": ["point 1", "point 2", "point 3"],
      "examples": ["example 1", "example 2"],
      "formulas": ["formula 1", "formula 2"],
      "existing_wikilinks": ["Existing Concept 1", "Existing Concept 2"],
      "related_concepts": ["Related Concept 1", "Related Concept 2"],
      "source": "{file_name}",
      "notes": "Any important context"
    }}
  ]
}}

IMPORTANT:
- Extract EVERY concept, not just the main one
- Each concept should be specific (e.g. "Price Elasticity of Demand" not just "Elasticity")
- Include definitions, formulas, and examples for each
- Reference existing wiki concepts using their exact names for [[Wikilinks]]
"""

    try:
        response_text = call_gemini(prompt)
        if not response_text:
            print(f"   All models failed")
            return None

        # Find JSON in response
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None

        json_str = response_text[start_idx:end_idx]
        result = json.loads(json_str)
        return result.get("concepts", [])

    except json.JSONDecodeError as e:
        print(f"   JSON parsing error: {e}")
        print(f"Response text: {response_text[:500]}")
        return None
    except Exception as e:
        print(f"   LLM error: {e}")
        return None


def create_concept_markdown(concept_data, existing_concepts, course_name=None):
    """Create markdown file with proper [[Wikilinks]]. Text only, no images."""

    title = concept_data["primary_concept"]
    slug = make_slug(title)

    # Flatten tiers for wikilink validation
    flat_concepts = flatten_tiers(existing_concepts)

    # Build wikilinks section from existing concepts only
    wikilinks = []
    for concept in concept_data.get("existing_wikilinks", []):
        if concept in flat_concepts:
            wikilinks.append(f"- [[{concept}]]")

    for concept in concept_data.get("related_concepts", []):
        if concept in flat_concepts and concept not in concept_data.get("existing_wikilinks", []):
            wikilinks.append(f"- [[{concept}]]")

    wikilinks_section = "\n".join(wikilinks) if wikilinks else "(No related concepts in wiki yet)"

    # Build sections
    formulas = concept_data.get("formulas", [])
    formulas_section = "\n".join(f"- {f}" for f in formulas) if formulas else "(No formulas)"

    examples = concept_data.get("examples", [])
    examples_section = "\n".join(f"- {e}" for e in examples) if examples else "(No examples)"

    key_points = concept_data.get("key_points", [])
    key_points_section = "\n".join(f"- {p}" for p in key_points) if key_points else "(No key points)"

    source = concept_data.get("source", "Lecture material")

    # Course metadata line
    course_line = f"\n**Course:** {course_name}" if course_name else ""

    # Reference with course prefix
    if course_name:
        reference = f"- {course_name}: {source}"
    else:
        reference = f"- {source}"

    markdown = f"""# {title}
{course_line}
**Source:** {source}

## Definition

{concept_data.get("definition", "(No definition extracted)")}

## Key Points

{key_points_section}

## Formulas & Equations

{formulas_section}

## Examples

{examples_section}

## Related Concepts

{wikilinks_section}

## Notes

{concept_data.get("notes", "Extracted from lecture material.")}

## References

{reference}
"""

    return slug, markdown


def merge_concepts_with_llm(merge_requests, existing_concepts, course_name=None):
    """
    Batch merge: send all duplicate concepts to Gemini in 1 API call.
    Each request has existing markdown + new data.
    Returns rewritten markdown for each concept.
    """
    if not merge_requests:
        return {}

    # Build the prompt with all concepts to merge
    merge_items = []
    for slug, data in merge_requests.items():
        merge_items.append(f"""
--- CONCEPT TO MERGE: {slug} ---
EXISTING MARKDOWN:
{data['existing_markdown']}

NEW CONTENT TO ADD:
Title: {data['new_data']['primary_concept']}
Definition: {data['new_data'].get('definition', '')}
Key Points: {json.dumps(data['new_data'].get('key_points', []))}
Examples: {json.dumps(data['new_data'].get('examples', []))}
Formulas: {json.dumps(data['new_data'].get('formulas', []))}
Source: {data['new_data'].get('source', '')}
Notes: {data['new_data'].get('notes', '')}
Course: {course_name or 'Unknown'}
--- END {slug} ---
""")

    flat_concepts = flatten_tiers(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in flat_concepts.keys()) if flat_concepts else "(none)"

    course_merge_instruction = ""
    if course_name:
        course_merge_instruction = f"""- The **Course:** metadata line should list ALL courses this concept appears in, comma-separated (e.g. "**Course:** Microeconomics, Marketing")
- If the existing page already has a **Course:** line, append "{course_name}" if not already listed
- If no **Course:** line exists, add "**Course:** {course_name}" after the title
- References should be prefixed with course name (e.g. "- {course_name}: filename.pdf")"""

    prompt = f"""You are merging new lecture content into existing wiki concept pages.

For each concept below, rewrite the FULL markdown page that seamlessly integrates the new content.

RULES:
- Merge new key points, examples, formulas into the existing sections (don't create "Additional" subsections)
- Remove duplicates - if a point already exists, don't add it again
- Keep the same markdown structure: # Title, **Course:**, **Source:**, ## Definition, ## Key Points, ## Formulas & Equations, ## Examples, ## Related Concepts, ## Notes, ## References
- Add ALL sources to a ## References section at the bottom (list each source file)
- Use [[Wikilinks]] only for these existing concepts: {existing_list}
- Keep content comprehensive but not repetitive
{course_merge_instruction}

CONCEPTS TO MERGE:
{"".join(merge_items)}

RESPOND ONLY WITH JSON (no markdown wrapping):
{{
  "merged": {{
    "{list(merge_requests.keys())[0]}": "full rewritten markdown here...",
    ...
  }}
}}
"""

    try:
        response_text = call_gemini(prompt)
        if not response_text:
            print(f"   All models failed for merge")
            return {}

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        result = json.loads(response_text[start_idx:end_idx])
        return result.get('merged', {})

    except Exception as e:
        print(f"   Merge API error: {e}")
        return {}


def _fuzzy_match(concept_title, concepts_dict):
    """Fuzzy match a title against a dict of {title: slug}. Returns slug or None."""
    lower_title = concept_title.lower()
    for existing_title, slug in concepts_dict.items():
        existing_lower = existing_title.lower()
        if (lower_title in existing_lower or
            existing_lower in lower_title or
            len(set(lower_title.split()) & set(existing_lower.split())) >= 2):
            return slug
    return None


def check_for_duplicates(concept_title, existing_concepts):
    """Check if concept already exists using tiered matching.

    Accepts either a flat dict {title: slug} (legacy) or a tiered dict
    {"same_course": {}, "same_group": {}, "other": {}}.

    Tiered logic:
    - same_course: exact + fuzzy
    - same_group: exact only
    - other: exact only
    """
    # Legacy flat dict mode
    if "same_course" not in existing_concepts:
        if concept_title in existing_concepts:
            return existing_concepts[concept_title]
        return _fuzzy_match(concept_title, existing_concepts)

    # Tiered mode
    same_course = existing_concepts["same_course"]
    same_group = existing_concepts["same_group"]
    other = existing_concepts["other"]

    # 1. Same course: exact + fuzzy
    if concept_title in same_course:
        return same_course[concept_title]
    fuzzy = _fuzzy_match(concept_title, same_course)
    if fuzzy:
        return fuzzy

    # 2. Same group: exact only
    if concept_title in same_group:
        return same_group[concept_title]

    # 3. Other: exact only
    if concept_title in other:
        return other[concept_title]

    return None


def flatten_tiers(tiered):
    """Flatten a tiered dict into a single {title: slug} dict."""
    if "same_course" not in tiered:
        return tiered  # already flat
    flat = {}
    flat.update(tiered["other"])
    flat.update(tiered["same_group"])
    flat.update(tiered["same_course"])
    return flat


def scoped_concepts_for_prompt(tiered):
    """Return the concept dict to embed in LLM prompts (same_course + same_group only).
    Falls back to flat dict if not tiered.
    """
    if "same_course" not in tiered:
        return tiered
    scoped = {}
    scoped.update(tiered["same_group"])
    scoped.update(tiered["same_course"])
    return scoped


def same_course_only(tiered):
    """Return just the same_course tier. Falls back to flat dict if not tiered."""
    if "same_course" not in tiered:
        return tiered
    return tiered["same_course"]


# ── Case Study Processing ────────────────────────────────────────────────────

def extract_case_with_llm(file_content, file_name, existing_concepts):
    """
    Single Gemini API call to extract case study data.
    Returns case study JSON data.
    """
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    existing_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"

    prompt = f"""You are extracting a business case study for a wiki.

FILE: {file_name}

EXISTING WIKI CONCEPTS (reference these using [[Wikilinks]]):
{existing_list}

TASK:
Read the provided business case study or simulation brief. Extract the critical data required for a strategic debrief.

FULL CASE CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "case_name": "The formal case name (e.g. 'Heidi Roizen at KPCB')",
  "core_dilemma": "Two to three sentences defining the primary strategic, leadership, or operational problem the protagonist is facing.",
  "stakeholders": [
    {{
      "name": "Person/Entity Name",
      "role": "Their role in the case",
      "incentives": "Their primary motivations or pressure points"
    }}
  ],
  "financial_context": "Any relevant margins, valuation multiples, budgets, or macro-economic constraints mentioned. If none, say 'No specific financial data provided.'",
  "related_concepts": ["Existing Concept 1", "Existing Concept 2"],
  "course": "",
  "source": "{file_name}"
}}

IMPORTANT:
- The case_name should be the formal title of the case, not a generic description
- Stakeholders should include ALL key actors, their roles, and motivations
- Financial context should include specific numbers when available
- related_concepts should ONLY reference concepts from the existing wiki list above
"""

    try:
        response_text = call_gemini(prompt)
        if not response_text:
            print(f"   All models failed")
            return None

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None

        json_str = response_text[start_idx:end_idx]
        result = json.loads(json_str)
        return result

    except json.JSONDecodeError as e:
        print(f"   JSON parsing error: {e}")
        return None
    except Exception as e:
        print(f"   LLM error: {e}")
        return None


def create_case_markdown(case_data, existing_concepts, course_name=None):
    """Create Case-*.md file with schema.md sections."""
    case_name = case_data["case_name"]
    slug = make_slug(case_name)

    # Flatten tiers for wikilink validation
    flat_concepts = flatten_tiers(existing_concepts)

    # Build stakeholders section
    stakeholders = case_data.get("stakeholders", [])
    if stakeholders:
        stakeholders_section = "\n".join(
            f"- **{s['name']}** ({s.get('role', 'Unknown role')}): {s.get('incentives', '')}"
            for s in stakeholders
        )
    else:
        stakeholders_section = "(No stakeholders identified)"

    # Build related concepts wikilinks (only existing ones)
    wikilinks = []
    for concept in case_data.get("related_concepts", []):
        if concept in flat_concepts:
            wikilinks.append(f"- [[{concept}]]")
    wikilinks_section = "\n".join(wikilinks) if wikilinks else "(No related concepts in wiki yet)"

    source = case_data.get("source", "Case material")
    course_line = f"\n**Course:** {course_name}" if course_name else ""

    if course_name:
        reference = f"- {course_name}: {source}"
    else:
        reference = f"- {source}"

    markdown = f"""# Case: {case_name}
**Tags:** #case-study #unresolved
{course_line}
**Source:** {source}

## 1. Core Dilemma

{case_data.get("core_dilemma", "(No dilemma extracted)")}

## 2. Key Stakeholders & Incentives

{stakeholders_section}

## 3. Financial Context & Constraints

{case_data.get("financial_context", "(No financial context extracted)")}

## 4. Class Discussion & Takeaways

*(To be populated by transcript processing)*

## Related Concepts

{wikilinks_section}

## References

{reference}
"""

    return slug, markdown


def check_case_duplicates(case_name, existing_cases):
    """Check if case already exists (exact or similar).
    Supports both flat dict and tiered dict formats.
    """
    flat = flatten_tiers(existing_cases)

    # Exact match
    if case_name in flat:
        return flat[case_name]

    # Fuzzy match (case-insensitive)
    lower_name = case_name.lower()
    for existing_name, slug in flat.items():
        existing_lower = existing_name.lower()
        if (lower_name in existing_lower or
            existing_lower in lower_name or
            len(set(lower_name.split()) & set(existing_lower.split())) >= 2):
            return slug

    return None


# ── Transcript Processing ────────────────────────────────────────────────────

def extract_transcript_with_llm(file_content, file_name, existing_concepts, existing_cases):
    """
    Single Gemini API call to extract BOTH concepts and case discussions from a transcript.
    Returns dual JSON: concepts[] + case_discussions[].
    """
    prompt_concepts = scoped_concepts_for_prompt(existing_concepts)
    prompt_cases = same_course_only(existing_cases)
    existing_concept_list = "\n".join(f"- {title}" for title in prompt_concepts.keys()) if prompt_concepts else "(none yet)"
    existing_case_list = "\n".join(f"- {title}" for title in prompt_cases.keys()) if prompt_cases else "(none yet)"

    prompt = f"""You are processing a class transcript for a wiki. This transcript may contain:
1. New concepts being taught (same as a lecture)
2. Discussion of case studies that were previously assigned

FILE: {file_name}

EXISTING WIKI CONCEPTS (use these exact names for [[Wikilinks]]):
{existing_concept_list}

EXISTING CASE STUDIES (match discussions to these case names):
{existing_case_list}

TASK:
Extract BOTH concepts AND case study discussions from this transcript.

FULL TRANSCRIPT CONTENT:
{file_content}

RESPOND ONLY WITH JSON (no markdown, no explanation):
{{
  "concepts": [
    {{
      "primary_concept": "Concept Title",
      "definition": "Clear definition in 2-3 sentences",
      "key_points": ["point 1", "point 2", "point 3"],
      "examples": ["example 1", "example 2"],
      "formulas": ["formula 1", "formula 2"],
      "existing_wikilinks": ["Existing Concept 1", "Existing Concept 2"],
      "related_concepts": ["Related Concept 1", "Related Concept 2"],
      "source": "{file_name}",
      "notes": "Any important context"
    }}
  ],
  "case_discussions": [
    {{
      "case_name": "Exact Case Name from existing cases list",
      "discussion_summary": "2-3 paragraph summary of how this case was discussed in class",
      "key_takeaways": ["takeaway 1", "takeaway 2"],
      "student_perspectives": ["perspective 1", "perspective 2"],
      "professor_insights": ["insight 1", "insight 2"]
    }}
  ]
}}

IMPORTANT:
- Extract ALL concepts taught, just like a lecture
- For case_discussions, ONLY include cases that were actually discussed in this transcript
- case_name must match an existing case from the list above (use the exact name)
- If no cases were discussed, return an empty case_discussions array
- student_perspectives: notable student arguments, debates, or questions
- professor_insights: key points the professor emphasized or corrected
"""

    try:
        response_text = call_gemini(prompt)
        if not response_text:
            print(f"   All models failed")
            return None, None

        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1

        if start_idx == -1 or end_idx <= start_idx:
            print(f"   LLM response doesn't contain JSON")
            print(f"Response: {response_text[:500]}")
            return None, None

        json_str = response_text[start_idx:end_idx]
        result = json.loads(json_str)
        return result.get("concepts", []), result.get("case_discussions", [])

    except json.JSONDecodeError as e:
        print(f"   JSON parsing error: {e}")
        return None, None
    except Exception as e:
        print(f"   LLM error: {e}")
        return None, None


def update_case_discussion(case_slug, discussion_data, source_file):
    """
    Update a Case-*.md file's '## 4. Class Discussion & Takeaways' section.
    If placeholder text exists, replace it. If content already exists, append with separator.
    Changes tag from #unresolved to #resolved.
    """
    case_file = WIKI_DIR / f"{CASE_FILE_PREFIX}{case_slug}{CONCEPT_FILE_SUFFIX}"

    if not case_file.exists():
        print(f"   Case file not found: {case_file}")
        return False

    try:
        with open(case_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Build discussion content
        discussion_md = discussion_data.get("discussion_summary", "")

        takeaways = discussion_data.get("key_takeaways", [])
        if takeaways:
            discussion_md += "\n\n### Key Takeaways\n\n"
            discussion_md += "\n".join(f"- {t}" for t in takeaways)

        student_perspectives = discussion_data.get("student_perspectives", [])
        if student_perspectives:
            discussion_md += "\n\n### Student Perspectives\n\n"
            discussion_md += "\n".join(f"- {p}" for p in student_perspectives)

        professor_insights = discussion_data.get("professor_insights", [])
        if professor_insights:
            discussion_md += "\n\n### Professor Insights\n\n"
            discussion_md += "\n".join(f"- {i}" for i in professor_insights)

        discussion_md += f"\n\n*Source: {source_file}*"

        # Find the section and replace/append
        placeholder = "*(To be populated by transcript processing)*"
        section_header = "## 4. Class Discussion & Takeaways"

        if placeholder in content:
            # Replace placeholder with actual content
            content = content.replace(placeholder, discussion_md)
        elif section_header in content:
            # Section exists with content — append with separator
            # Find the next ## section after the discussion section
            section_start = content.index(section_header) + len(section_header)
            next_section = content.find("\n## ", section_start)

            if next_section != -1:
                # Insert before the next section
                existing_content = content[section_start:next_section].rstrip()
                content = (content[:section_start] +
                          existing_content + "\n\n---\n\n" + discussion_md + "\n\n" +
                          content[next_section:])
            else:
                # No next section — append at end
                content = content.rstrip() + "\n\n---\n\n" + discussion_md + "\n"

        # Change tag from #unresolved to #resolved
        content = content.replace("#unresolved", "#resolved")

        with open(case_file, 'w', encoding='utf-8') as f:
            f.write(content)

        return True

    except Exception as e:
        print(f"   Error updating case discussion: {e}")
        return False


# ── Main Processing Flows ────────────────────────────────────────────────────

def process_lecture_file(content, file_path, existing_concepts, course_name):
    """Process a lecture file: extract concepts, create/merge."""
    api_calls = 0

    # Extract concepts (1 API call)
    print(f"\n   Using Gemini to extract concepts (1 API call)...")
    concepts_list = extract_concepts_with_llm(content, Path(file_path).name, existing_concepts)
    api_calls += 1

    if not concepts_list:
        print("   Failed to extract concept data")
        return api_calls

    print(f"   Extracted {len(concepts_list)} concepts")

    # Display results
    print(f"\n{'='*70}")
    print(f"   Extraction Results:")
    print(f"{'='*70}")
    for i, concept_data in enumerate(concepts_list, 1):
        print(f"\n  {i}. {concept_data['primary_concept']}")
        print(f"     Definition: {concept_data.get('definition', '')[:80]}...")
        print(f"     Key Points: {len(concept_data.get('key_points', []))} | Examples: {len(concept_data.get('examples', []))} | Formulas: {len(concept_data.get('formulas', []))}")

    # Separate new concepts from duplicates
    created = 0
    updated = 0
    merge_requests = {}

    for concept_data in concepts_list:
        title = concept_data["primary_concept"]
        print(f"\n{'~'*50}")
        print(f"   Processing: {title}")

        duplicate = check_for_duplicates(title, existing_concepts)

        if duplicate:
            print(f"   Duplicate found -> queuing for merge (Concept-{duplicate}.md)")
            existing_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{duplicate}{CONCEPT_FILE_SUFFIX}"

            try:
                with open(existing_file, 'r', encoding='utf-8') as f:
                    existing_markdown = f.read()

                if duplicate in merge_requests:
                    prev = merge_requests[duplicate]['new_data']
                    prev['key_points'] = prev.get('key_points', []) + concept_data.get('key_points', [])
                    prev['examples'] = prev.get('examples', []) + concept_data.get('examples', [])
                    prev['formulas'] = prev.get('formulas', []) + concept_data.get('formulas', [])
                else:
                    merge_requests[duplicate] = {
                        'existing_markdown': existing_markdown,
                        'new_data': concept_data
                    }
            except Exception as e:
                print(f"   Error reading existing file: {e}")
        else:
            slug, markdown = create_concept_markdown(concept_data, existing_concepts, course_name)
            output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(markdown)
                print(f"   Created: {output_file}")
                created += 1
                # Register in same_course tier if tiered, else flat
                if "same_course" in existing_concepts:
                    existing_concepts["same_course"][title] = slug
                else:
                    existing_concepts[title] = slug
            except Exception as e:
                print(f"   Error saving: {e}")

    # Batch merge duplicates (1 API call if needed)
    if merge_requests:
        print(f"\n{'~'*50}")
        print(f"   Merging {len(merge_requests)} existing concepts with new content (1 API call)...")
        api_calls += 1

        merged_results = merge_concepts_with_llm(merge_requests, existing_concepts, course_name)

        for slug, new_markdown in merged_results.items():
            output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(new_markdown)
                print(f"   Merged: {output_file}")
                updated += 1
            except Exception as e:
                print(f"   Error saving merge: {e}")

    print(f"\n   Created: {created} new concepts")
    print(f"   Updated: {updated} existing concepts")
    return api_calls


def process_case_file(content, file_path, existing_concepts, existing_cases, course_name):
    """Process a case study file: extract case data, create Case-*.md."""
    api_calls = 0

    # Extract case data (1 API call)
    print(f"\n   Using Gemini to extract case study (1 API call)...")
    case_data = extract_case_with_llm(content, Path(file_path).name, existing_concepts)
    api_calls += 1

    if not case_data:
        print("   Failed to extract case data")
        return api_calls

    case_name = case_data.get("case_name", "Unknown Case")
    print(f"   Extracted case: {case_name}")

    # Display results
    print(f"\n{'='*70}")
    print(f"   Case Study Results:")
    print(f"{'='*70}")
    print(f"   Name: {case_name}")
    print(f"   Dilemma: {case_data.get('core_dilemma', '')[:100]}...")
    print(f"   Stakeholders: {len(case_data.get('stakeholders', []))}")
    print(f"   Related concepts: {len(case_data.get('related_concepts', []))}")

    # Check for duplicate case
    duplicate = check_case_duplicates(case_name, existing_cases)

    if duplicate:
        print(f"\n   Case already exists: Case-{duplicate}.md")
        print(f"   Skipping (case files are not auto-merged)")
    else:
        slug, markdown = create_case_markdown(case_data, existing_concepts, course_name)
        output_file = WIKI_DIR / f"{CASE_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown)
            print(f"\n   Created: {output_file}")
            if "same_course" in existing_cases:
                existing_cases["same_course"][case_name] = slug
            else:
                existing_cases[case_name] = slug
        except Exception as e:
            print(f"   Error saving case: {e}")

    return api_calls


def process_transcript_file(content, file_path, existing_concepts, existing_cases, course_name):
    """Process a transcript: extract concepts AND update case discussions."""
    api_calls = 0

    # Extract both concepts and case discussions (1 API call)
    print(f"\n   Using Gemini to extract concepts + case discussions (1 API call)...")
    concepts_list, case_discussions = extract_transcript_with_llm(
        content, Path(file_path).name, existing_concepts, existing_cases
    )
    api_calls += 1

    if concepts_list is None:
        print("   Failed to extract transcript data")
        return api_calls

    print(f"   Extracted {len(concepts_list)} concepts, {len(case_discussions or [])} case discussions")

    # ── Part 1: Process concepts (reuse lecture logic) ──
    if concepts_list:
        print(f"\n{'='*70}")
        print(f"   Concept Extraction Results:")
        print(f"{'='*70}")
        for i, concept_data in enumerate(concepts_list, 1):
            print(f"\n  {i}. {concept_data['primary_concept']}")
            print(f"     Definition: {concept_data.get('definition', '')[:80]}...")

        created = 0
        updated = 0
        merge_requests = {}

        for concept_data in concepts_list:
            title = concept_data["primary_concept"]
            print(f"\n{'~'*50}")
            print(f"   Processing: {title}")

            duplicate = check_for_duplicates(title, existing_concepts)

            if duplicate:
                print(f"   Duplicate found -> queuing for merge (Concept-{duplicate}.md)")
                existing_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{duplicate}{CONCEPT_FILE_SUFFIX}"

                try:
                    with open(existing_file, 'r', encoding='utf-8') as f:
                        existing_markdown = f.read()

                    if duplicate in merge_requests:
                        prev = merge_requests[duplicate]['new_data']
                        prev['key_points'] = prev.get('key_points', []) + concept_data.get('key_points', [])
                        prev['examples'] = prev.get('examples', []) + concept_data.get('examples', [])
                        prev['formulas'] = prev.get('formulas', []) + concept_data.get('formulas', [])
                    else:
                        merge_requests[duplicate] = {
                            'existing_markdown': existing_markdown,
                            'new_data': concept_data
                        }
                except Exception as e:
                    print(f"   Error reading existing file: {e}")
            else:
                slug, markdown = create_concept_markdown(concept_data, existing_concepts, course_name)
                output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"

                try:
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(markdown)
                    print(f"   Created: {output_file}")
                    created += 1
                    if "same_course" in existing_concepts:
                        existing_concepts["same_course"][title] = slug
                    else:
                        existing_concepts[title] = slug
                except Exception as e:
                    print(f"   Error saving: {e}")

        # Batch merge duplicates (1 API call if needed)
        if merge_requests:
            print(f"\n{'~'*50}")
            print(f"   Merging {len(merge_requests)} existing concepts with new content (1 API call)...")
            api_calls += 1

            merged_results = merge_concepts_with_llm(merge_requests, existing_concepts, course_name)

            for slug, new_markdown in merged_results.items():
                output_file = WIKI_DIR / f"{CONCEPT_FILE_PREFIX}{slug}{CONCEPT_FILE_SUFFIX}"
                try:
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(new_markdown)
                    print(f"   Merged: {output_file}")
                    updated += 1
                except Exception as e:
                    print(f"   Error saving merge: {e}")

        print(f"\n   Concepts — Created: {created}, Updated: {updated}")

    # ── Part 2: Update case discussions (file I/O only, no API calls) ──
    if case_discussions:
        print(f"\n{'='*70}")
        print(f"   Case Discussion Updates:")
        print(f"{'='*70}")

        discussions_updated = 0
        for disc in case_discussions:
            case_name = disc.get("case_name", "")
            print(f"\n   Updating case: {case_name}")

            # Find matching case slug
            case_slug = flatten_tiers(existing_cases).get(case_name)
            if not case_slug:
                # Try fuzzy match
                case_slug = check_case_duplicates(case_name, existing_cases)

            if case_slug:
                success = update_case_discussion(case_slug, disc, Path(file_path).name)
                if success:
                    print(f"   Updated: Case-{case_slug}.md")
                    discussions_updated += 1
                else:
                    print(f"   Failed to update Case-{case_slug}.md")
            else:
                print(f"   No matching case found for: {case_name}")

        print(f"\n   Case discussions updated: {discussions_updated}")

    return api_calls


def main():
    """Main processing function."""

    parsed = parse_args()

    # Handle --seed mode (no file needed, no API calls)
    if parsed['seed']:
        course_name = parsed['course']
        if not course_name:
            print("\nError: --course is required for seeding.")
            print('Usage: python process_single_file.py --seed --course "CourseName"')
            sys.exit(1)
        print(f"\n{'='*70}")
        print(f"   Seed Mode: {course_name}")
        print(f"{'='*70}")
        created = seed_concepts(course_name)
        if created > 0:
            total = len(load_existing_concepts())
            print(f"\n   Total concepts in wiki: {total}")
            log_ingestion("seed", f"{course_name} seed concepts", "seed", f"{created} concepts created")
        return

    if not parsed['file_path']:
        print("\nUsage:")
        print('   python process_single_file.py <file_path> --course "CourseName"')
        print('   python process_single_file.py <file_path> --course "CourseName" --type case')
        print('   python process_single_file.py <file_path> --course "CourseName" --type transcript')
        print('   python process_single_file.py --seed --course "CourseName"')
        print("\nTypes: lecture (default), case, transcript")
        sys.exit(1)

    file_path = parsed['file_path']
    course_name = parsed['course']
    skip_images = parsed['no_images']
    file_type = parsed['type']

    if not course_name:
        print("\nError: --course is required.")
        print('Usage: python process_single_file.py <file_path> --course "CourseName"')
        sys.exit(1)

    print(f"\n{'='*70}")
    print(f"   Processing: {file_path}")
    print(f"   Course: {course_name}")
    print(f"   Type: {file_type}")
    print(f"{'='*70}")

    # Step 1: Setup
    if not setup_gemini():
        sys.exit(1)

    # Step 2: Load course groups and existing concepts/cases
    group_courses = load_course_groups(course_name)
    if group_courses:
        print(f"\n   Course group siblings: {', '.join(group_courses)}")
    else:
        print(f"\n   No course group found for '{course_name}' (using exact-only cross-course matching)")

    print("\n   Loading existing concepts...")
    existing_concepts = load_existing_concepts(course_name, group_courses)
    flat_concepts = flatten_tiers(existing_concepts)
    sc = len(existing_concepts.get("same_course", {}))
    sg = len(existing_concepts.get("same_group", {}))
    ot = len(existing_concepts.get("other", {}))
    print(f"   Found {len(flat_concepts)} existing concepts (same_course={sc}, same_group={sg}, other={ot})")

    existing_cases = load_existing_cases(course_name, group_courses)
    flat_cases = flatten_tiers(existing_cases)
    sc_c = len(existing_cases.get("same_course", {}))
    sg_c = len(existing_cases.get("same_group", {}))
    ot_c = len(existing_cases.get("other", {}))
    print(f"   Found {len(flat_cases)} existing cases (same_course={sc_c}, same_group={sg_c}, other={ot_c})")

    if flat_concepts:
        print("\n   Sample existing concepts:")
        for title in list(flat_concepts.keys())[:5]:
            print(f"      - {title}")
        if len(flat_concepts) > 5:
            print(f"      ... and {len(flat_concepts) - 5} more")

    if flat_cases:
        print("\n   Existing cases:")
        for title in flat_cases.keys():
            print(f"      - {title}")

    # Step 3: Extract images from PDF (optional, only for lectures/cases)
    if file_type in ('lecture', 'case') and Path(file_path).suffix.lower() == '.pdf' and not skip_images:
        print(f"\n   Extracting charts from PDF (saving to assets folder)...")
        extracted_images = extract_images_from_pdf(file_path, WIKI_DIR / "assets" / "charts")
        if extracted_images:
            print(f"   Saved {len(extracted_images)} charts to MBAWiki/assets/charts/")
        else:
            print(f"   No charts found in PDF")
    elif skip_images:
        print(f"\n   Image extraction: SKIPPED (--no-images)")

    # Step 4: Extract text
    print(f"\n   Extracting text from file...")
    try:
        content = extract_text_from_file(file_path)
        if not content:
            print("   No content extracted")
            sys.exit(1)
        print(f"   Read {len(content):,} characters")
    except FileNotFoundError as e:
        print(f"   {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"   {e}")
        sys.exit(1)

    # Step 5: Dispatch based on type
    if file_type == 'case':
        api_calls = process_case_file(content, file_path, existing_concepts, existing_cases, course_name)
    elif file_type == 'transcript':
        api_calls = process_transcript_file(content, file_path, existing_concepts, existing_cases, course_name)
    else:
        api_calls = process_lecture_file(content, file_path, existing_concepts, course_name)

    # Step 6: Log the ingestion
    source_name = Path(file_path).name
    log_details = f"{api_calls} API calls"
    log_ingestion("ingest", source_name, file_type, log_details)

    # Step 6b: Update search index incrementally
    try:
        from build_search_index import build_index
        total, updated = build_index(append_mode=True)
        print(f"   Search index: {total} entries ({updated} re-embedded) [OK]")
    except ImportError:
        print(f"   Search index: skipped (fastembed not installed - run: pip install fastembed)")
    except Exception as e:
        print(f"   Search index: failed (non-fatal): {e}")

    # Step 7: Summary
    print(f"\n{'='*70}")
    print(f"   Complete!")
    print(f"{'='*70}")
    total_concepts = len(load_existing_concepts())
    total_cases = len(load_existing_cases())
    print(f"   Total in wiki: {total_concepts} concepts, {total_cases} cases")
    print(f"   API calls used: {api_calls}")
    print(f"   Logged to: log.md ✓")

    print(f"\n   Next steps:")
    print(f"   1. Review the markdown files in MBAWiki/")
    print(f"   2. Restart the wiki server:")
    print(f"      python wiki_viewer/app.py")
    print(f"   3. View in browser:")
    print(f"      http://127.0.0.1:5000/")
    print(f"\n")


if __name__ == "__main__":
    main()
